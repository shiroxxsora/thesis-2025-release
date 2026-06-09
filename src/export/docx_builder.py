"""
Генерация DOCX-документов по отдельным участкам.
"""

import logging
import os
from typing import Optional

import pandas as pd

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

from .coordinate_presenter import compute_distances

logger = logging.getLogger(__name__)


class DOCXBuilder:
    """Построитель DOCX-документов для отдельных участков."""
    
    def __init__(self):
        """Инициализация построителя DOCX."""
        if not DOCX_AVAILABLE:
            raise RuntimeError(
                "Требуется пакет 'python-docx' (pip install python-docx) для экспорта DOCX"
            )
    
    def generate(
        self,
        output_path: str,
        cadastral_number: str,
        cadastral_df: pd.DataFrame,
        coords_df: pd.DataFrame,
        violations_df: pd.DataFrame,
        viol_coords_df: pd.DataFrame,
        image_path: Optional[str] = None,
        overview_path: Optional[str] = None,
    ) -> str:
        """
        Генерирует DOCX-документ для участка.
        
        Args:
            output_path: Путь для сохранения DOCX
            cadastral_number: Кадастровый номер участка
            cadastral_df: DataFrame с кадастровыми участками
            coords_df: DataFrame с координатами участков
            violations_df: DataFrame с нарушениями
            viol_coords_df: DataFrame с координатами нарушений
            image_path: Путь к zoom-карте (опционально)
            overview_path: Путь к обзорной карте (опционально)
            
        Returns:
            Путь к созданному DOCX-файлу
        """
        doc = Document()
        
        # Настройка стилей для кириллицы
        self._setup_styles(doc)
        
        # Заголовок
        h = doc.add_paragraph('СХЕМА')
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = doc.add_paragraph('обмера земельного участка')
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Данные по участку
        doc.add_paragraph(f"Кадастровый номер: {cadastral_number}")
        
        # Сводка по нарушениям
        v = violations_df[violations_df['Ближайший кадастровый номер'] == cadastral_number]
        if not v.empty:
            self._add_violations_summary(doc, v, viol_coords_df)
        
        # Обзорная карта
        if overview_path and os.path.exists(overview_path):
            try:
                doc.add_paragraph('Обзорная карта участка')
                doc.add_picture(overview_path, width=Inches(5.5))
            except Exception:
                pass
        
        # Детальная карта
        if image_path and os.path.exists(image_path):
            try:
                doc.add_paragraph('Детальная карта нарушений')
                doc.add_picture(image_path, width=Inches(5.5))
            except Exception:
                pass
        
        doc.save(output_path)
        logger.info(f"DOCX создан: {output_path}")
        return output_path
    
    def _setup_styles(self, doc):
        """Настраивает стили документа для поддержки кириллицы."""
        for style_name in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3']:
            try:
                st = doc.styles[style_name]
                st.font.name = 'Times New Roman'
                st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                st.font.size = Pt(11)
            except Exception:
                pass
    
    def _add_violations_summary(
        self,
        doc,
        violations_df: pd.DataFrame,
        viol_coords_df: pd.DataFrame,
    ):
        """Добавляет сводку по нарушениям (координаты как в Excel)."""
        sum_violation = float(violations_df['Площадь нарушения, м²'].fillna(0).sum())
        doc.add_paragraph(f"Количество нарушений: {len(violations_df)}")
        doc.add_paragraph(
            f"Площадь нарушений по участку (объединение контуров): {sum_violation:.2f} м²"
        )
        
        # Краткая таблица нарушений
        v_table = doc.add_table(rows=1, cols=4)
        try:
            v_table.style = 'Table Grid'
        except Exception:
            pass
        
        v_hdr = v_table.rows[0].cells
        v_hdr[0].text = '№'
        v_hdr[1].text = 'Площадь, м²'
        v_hdr[2].text = 'Центроид X'
        v_hdr[3].text = 'Центроид Y'
        
        for i, (_, r) in enumerate(violations_df.iterrows(), 1):
            try:
                area_val = float(pd.to_numeric(r.get('Площадь нарушения, м²', 0), errors='coerce'))
            except Exception:
                area_val = 0
            try:
                cx = float(pd.to_numeric(r.get('Центроид X', 0), errors='coerce'))
            except Exception:
                cx = 0
            try:
                cy = float(pd.to_numeric(r.get('Центроид Y', 0), errors='coerce'))
            except Exception:
                cy = 0
            
            row = v_table.add_row().cells
            row[0].text = str(i)
            row[1].text = f"{area_val:.2f}"
            row[2].text = f"{cx:.2f}"
            row[3].text = f"{cy:.2f}"
        
        # Таблицы координат по каждому нарушению
        if not viol_coords_df.empty and '№ нарушения' in viol_coords_df.columns:
            self._add_violation_coords_tables(doc, violations_df, viol_coords_df)
    
    def _add_violation_coords_tables(
        self,
        doc,
        violations_df: pd.DataFrame,
        viol_coords_df: pd.DataFrame,
    ):
        """Добавляет таблицы координат (как в Excel)."""
        v_reset = violations_df.reset_index(drop=True)
        for local_idx, (idx_row, row_v) in enumerate(v_reset.iterrows(), 1):
            global_v_num = row_v.get('№ нарушения', idx_row + 1)
            try:
                global_v_num = int(global_v_num)
            except Exception:
                global_v_num = idx_row + 1
            
            sub = viol_coords_df[viol_coords_df['№ нарушения'] == global_v_num].copy()
            if sub.empty:
                continue
            
            sub.sort_values(by='Номер точки', inplace=True)
            pts = list(zip(
                pd.to_numeric(sub['X'], errors='coerce').astype(float),
                pd.to_numeric(sub['Y'], errors='coerce').astype(float)
            ))
            if len(pts) >= 2 and pts[0] == pts[-1]:
                pts = pts[:-1]
            
            logger.debug(f"DOCX: Нарушение №{local_idx} (глобальное #{global_v_num}): {len(pts)} точек")
            dists = compute_distances(pts)
            
            doc.add_paragraph(f"Координаты нарушения № {local_idx}")
            vt = doc.add_table(rows=1, cols=4)
            try:
                vt.style = 'Table Grid'
            except Exception:
                pass
            
            hdr = vt.rows[0].cells
            hdr[0].text = 'Обозначение\nточки'
            hdr[1].text = 'X'
            hdr[2].text = 'Y'
            hdr[3].text = 'Расстояние до точки, м.'
            
            for j, (x, y) in enumerate(pts, 1):
                dist_idx = (j - 2) % len(dists)
                px, py = float(x), float(y)
                rw = vt.add_row().cells
                rw[0].text = str(j)
                rw[1].text = f"{px:.2f}"
                rw[2].text = f"{py:.2f}"
                rw[3].text = f"{dists[dist_idx]:.2f}"
