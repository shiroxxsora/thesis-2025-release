import os
import sys
import math
import argparse
import warnings
from typing import List, Tuple, Optional, Dict, Any

import pandas as pd

# headless backend for matplotlib
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import Polygon as MplPolygon
from matplotlib.path import Path as MplPath
from matplotlib.patches import PathPatch
import numpy as np
from concurrent.futures import ThreadPoolExecutor, as_completed
from io import BytesIO

# Pillow для контроля размера изображений
try:
    from PIL import Image as PILImage
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from shapely import wkt as shapely_wkt
from shapely.geometry import Polygon, MultiPolygon, box
from shapely.ops import transform

# Зависимость для перепроецирования
try:
    import pyproj
    PYPROJ_AVAILABLE = True
except ImportError:
    PYPROJ_AVAILABLE = False

# GDAL для чтения GeoTIFF (как в других скриптах)
try:
    from osgeo import gdal, osr
    GDAL_AVAILABLE = True
except Exception:
    GDAL_AVAILABLE = False


# DOCX (опционально)
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


def _safe_filename(name: str) -> str:
    """Приводит строку к безопасному имени файла."""
    invalid = '\\/:*?"<>|'  # Windows-символы
    result = name
    for ch in invalid:
        result = result.replace(ch, '_')
    return result.replace(' ', '_')


def _format_float(value: float, ndigits: int = 2) -> str:
    try:
        return f"{float(value):.{ndigits}f}"
    except Exception:
        return ""


def _present_xy(x: float, y: float) -> Tuple[float, float]:
    try:
        xf = float(x)
        yf = float(y)
    except Exception:
        return x, y
    new_x = yf
    sign = -1.0 if xf < 0 else 1.0
    new_y = sign * (4000000.0 + abs(xf))
    return new_x, new_y


def compute_distances(points_xy: List[Tuple[float, float]]) -> List[float]:
    """
    Возвращает список расстояний от каждой точки до следующей по порядку
    (замыкаем на первую точку). Единицы соответствуют CRS (ожидаются метры).
    """
    if not points_xy:
        return []
    distances: List[float] = []
    for i in range(len(points_xy)):
        x1, y1 = points_xy[i]
        x2, y2 = points_xy[(i + 1) % len(points_xy)]
        distances.append(math.hypot(x2 - x1, y2 - y1))
    return distances


def _simplify_points(points: List[Tuple[float, float]], 
                     min_spacing: float = 3.0, 
                     max_points: Optional[int] = None) -> List[Tuple[float, float]]:
    if len(points) <= 3:
        return points
    
    simplified = [points[0]]
    
    for i in range(1, len(points)):
        x_prev, y_prev = simplified[-1]
        x_curr, y_curr = points[i]
        dist = math.hypot(x_curr - x_prev, y_curr - y_prev)
        
        if dist >= min_spacing:
            simplified.append(points[i])
    
    if max_points and len(simplified) > max_points:
        step = len(simplified) / max_points
        indices = [int(i * step) for i in range(max_points)]
        simplified = [simplified[i] for i in indices]
    
    return simplified


def load_report_frames(report_path: str) -> Dict[str, pd.DataFrame]:
    """Читает необходимые листы из Excel-отчета."""
    print(f"[INFO] Чтение отчёта: {report_path}", flush=True)
    xls = pd.ExcelFile(report_path)
    print(f"[INFO] Найдены листы: {xls.sheet_names}", flush=True)

    # Листы по именам из export_to_excel()
    try:
        cadastral_df = pd.read_excel(xls, sheet_name='2. Кадастровые участки')
        print(f"[INFO] '2. Кадастровые участки': {len(cadastral_df)} строк", flush=True)
    except Exception as e:
        print(f"[ERROR] Не удалось прочитать лист '2. Кадастровые участки': {e}", flush=True)
        raise

    try:
        violations_df = pd.read_excel(xls, sheet_name='3. Нарушения')
        print(f"[INFO] '3. Нарушения': {len(violations_df)} строк", flush=True)
    except Exception as e:
        print(f"[ERROR] Не удалось прочитать лист '3. Нарушения': {e}", flush=True)
        raise

    try:
        coords_df = pd.read_excel(xls, sheet_name='4. Координаты участков')
        print(f"[INFO] '4. Координаты участков': {len(coords_df)} строк", flush=True)
    except Exception as e:
        print(f"[ERROR] Не удалось прочитать лист '4. Координаты участков': {e}", flush=True)
        raise

    # Лист 5: Координаты нарушений (для нумерации точек и таблиц по нарушениям)
    try:
        viol_coords_df = pd.read_excel(xls, sheet_name='5. Координаты нарушений')
        print(f"[INFO] '5. Координаты нарушений': {len(viol_coords_df)} строк", flush=True)
    except Exception as e:
        print(f"[ERROR] Не удалось прочитать лист '5. Координаты нарушений': {e}", flush=True)
        raise

    return {
        'cadastral_df': cadastral_df,
        'violations_df': violations_df,
        'coords_df': coords_df,
        'viol_coords_df': viol_coords_df,
    }


def build_parcel_story(cadastral_number: str,
                       cadastral_df: pd.DataFrame,
                       coords_df: pd.DataFrame,
                       violations_df: pd.DataFrame,
                       viol_coords_df: pd.DataFrame,
                       image_path: Optional[str] = None,
                       overview_path: Optional[str] = None,
                       font_name: str = 'Helvetica') -> List:
    """
    Формирует содержимое PDF: заголовок, сведения по участку, таблицу координат,
    список и сумму нарушений, опционально изображение.
    """
    styles = getSampleStyleSheet()
    heading = ParagraphStyle(name='HeadingCenter', parent=styles['Heading2'], alignment=1, fontName=font_name)
    normal = ParagraphStyle(name='NormalCyr', parent=styles['Normal'], fontName=font_name)

    story: List = []

    story.append(Paragraph('СХЕМА', heading))
    story.append(Paragraph('обмера земельного участка', ParagraphStyle(name='Sub', parent=normal, alignment=1)))
    story.append(Spacer(1, 8))

    # Данные по участку
    story.append(Paragraph(f"Кадастровый номер: <b>{cadastral_number}</b>", normal))
    story.append(Spacer(1, 6))

    # Сводка по нарушениям для участка
    v = violations_df[violations_df['Ближайший кадастровый номер'] == cadastral_number]
    if not v.empty:
        sum_violation = float(v['Площадь нарушения, м²'].fillna(0).sum())
        story.append(Paragraph(f"Количество нарушений: <b>{len(v)}</b>", normal))
        story.append(Paragraph(f"Суммарная площадь нарушений: <b>{_format_float(sum_violation, 2)}</b> м²", normal))
        story.append(Spacer(1, 6))

        # Краткая таблица нарушений
        v_table = [['№', 'Площадь, м²', 'Центроид X', 'Центроид Y']]
        for i, (_, r) in enumerate(v.iterrows(), 1):
            try:
                area_val = pd.to_numeric(r.get('Площадь нарушения, м²', 0), errors='coerce')
            except Exception:
                area_val = 0
            try:
                cx = pd.to_numeric(r.get('Центроид X', 0), errors='coerce')
            except Exception:
                cx = 0
            try:
                cy = pd.to_numeric(r.get('Центроид Y', 0), errors='coerce')
            except Exception:
                cy = 0
            try:
                cx_p, cy_p = _present_xy(cx if pd.notna(cx) else 0, cy if pd.notna(cy) else 0)
            except Exception:
                cx_p, cy_p = cx, cy
            v_table.append([
                str(i),
                _format_float(area_val if pd.notna(area_val) else 0, 2),
                _format_float(cx_p, 2),
                _format_float(cy_p, 2),
            ])
        v_tbl = Table(v_table, repeatRows=1)
        v_tbl.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f7f7ff')),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('ALIGN', (0, 1), (0, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, -1), font_name),
        ]))
        story.append(v_tbl)
        story.append(Spacer(1, 8))

        # Таблицы координат по каждому нарушению (если есть лист 5)
        if not viol_coords_df.empty and '№ нарушения' in viol_coords_df.columns:
            v_reset = v.reset_index(drop=True)
            for local_idx, (idx_row, row_v) in enumerate(v_reset.iterrows(), 1):
                # Используем глобальный номер нарушения для поиска координат
                global_v_num = row_v['№ нарушения'] if '№ нарушения' in v.columns else (idx_row + 1)
                try:
                    global_v_num = int(global_v_num)
                except Exception:
                    global_v_num = idx_row + 1
                
                sub = viol_coords_df[viol_coords_df['№ нарушения'] == global_v_num].copy()
                if sub.empty:
                    continue
                try:
                    sub.sort_values(by='Номер точки', inplace=True)
                    pts = list(zip(pd.to_numeric(sub['X'], errors='coerce').astype(float),
                                    pd.to_numeric(sub['Y'], errors='coerce').astype(float)))
                    if len(pts) >= 2 and pts[0] == pts[-1]:
                        pts = pts[:-1]
                    print(f"[DEBUG] PDF: КН {cadastral_number}, Локальное нарушение №{local_idx} (глобальное #{global_v_num}): {len(pts)} точек в таблице.", flush=True)
                    dists = compute_distances(pts)
                    # Используем локальный номер в заголовке
                    story.append(Paragraph(f"Координаты нарушения № {local_idx}", heading))
                    vt = [['Обозначение\nточки', 'X', 'Y', 'Расстояние до точки, м.']]
                    for i, (x, y) in enumerate(pts, 1):
                        dist_idx = (i - 2) % len(dists)
                        px, py = _present_xy(x, y)
                        vt.append([str(i), _format_float(px, 2), _format_float(py, 2), _format_float(dists[dist_idx], 2)])
                    # без дополнительной дублирующей строки
                    vt_tbl = Table(vt, repeatRows=1)
                    vt_tbl.setStyle(TableStyle([
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#eef7ee')),
                        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                        ('ALIGN', (0, 1), (0, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, -1), font_name),
                    ]))
                    story.append(vt_tbl)
                    story.append(Spacer(1, 8))
                except Exception:
                    continue

    if overview_path and os.path.exists(overview_path):
        try:
            story.append(Paragraph('Обзорная карта участка', heading))
            story.append(Spacer(1, 6))
            img = Image(overview_path)
            max_w, max_h = 520, 520
            img._restrictSize(max_w, max_h)
            story.append(img)
            story.append(Spacer(1, 12))
        except Exception:
            pass

    if image_path and os.path.exists(image_path):
        try:
            story.append(Paragraph('Детальная карта нарушений', heading))
            story.append(Spacer(1, 6))
            img = Image(image_path)
            max_w, max_h = 520, 520
            img._restrictSize(max_w, max_h)
            story.append(img)
        except Exception:
            pass

    return story


def export_per_parcel_pdfs(report_path: str,
                           output_dir: str,
                           image_path: Optional[str] = None,
                           only_with_violations: bool = True,
                           limit: Optional[int] = None,
                           font_path: Optional[str] = None,
                           out_format: str = 'pdf',
                           proj_string: Optional[str] = None,
                           min_point_spacing: float = 3.0,
                           max_points: Optional[int] = None) -> None:
    print("[INFO] Старт экспорта по участкам", flush=True)
    frames = load_report_frames(report_path)
    cadastral_df = frames['cadastral_df']
    violations_df = frames['violations_df']
    coords_df = frames['coords_df']
    viol_coords_df = frames['viol_coords_df']

    # Регистрация шрифта с кириллицей
    font_name = 'Helvetica'
    font_candidates: List[str] = []
    if font_path:
        font_candidates.append(font_path)
    # типичные пути Windows/Linux
    font_candidates += [
        os.path.join('C:\\Windows\\Fonts', 'arial.ttf'),
        os.path.join('C:\\Windows\\Fonts', 'calibri.ttf'),
        os.path.join('C:\\Windows\\Fonts', 'times.ttf'),
        '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
        '/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf',
        '/usr/share/fonts/truetype/freefont/FreeSans.ttf',
    ]

    registered = False
    for candidate in font_candidates:
        if candidate and os.path.exists(candidate):
            try:
                pdfmetrics.registerFont(TTFont('CyrillicTT', candidate))
                font_name = 'CyrillicTT'
                print(f"[INFO] Использую шрифт: {candidate}", flush=True)
                registered = True
                break
            except Exception as e:
                print(f"[WARN] Не удалось зарегистрировать шрифт {candidate}: {e}", flush=True)
    if not registered:
        print("[WARN] Кириллический TTF не найден. Попробую стандартный шрифт (может не поддерживать кириллицу)", flush=True)

    if only_with_violations:
        parcel_ids = sorted(set(violations_df['Ближайший кадастровый номер'].dropna().astype(str)))
    else:
        parcel_ids = sorted(set(cadastral_df['Кадастровый номер'].dropna().astype(str)))

    if limit is not None:
        parcel_ids = parcel_ids[:max(0, int(limit))]

    print(f"[INFO] Всего участков к выгрузке: {len(parcel_ids)}", flush=True)
    os.makedirs(output_dir, exist_ok=True)

    # Предзагрузка подложки один раз
    bg_tuple: Optional[Tuple[Any, Optional[Tuple[float, float, float, float]], Any]] = None
    if image_path and os.path.exists(image_path):
        try:
            if image_path.lower().endswith(('.tif', '.tiff')):
                if not GDAL_AVAILABLE:
                    raise RuntimeError("Требуется GDAL (osgeo). Установите пакет 'gdal'/'osgeo'.")
                ds = gdal.Open(image_path)
                if ds is None:
                    raise RuntimeError(f"Не удалось открыть GeoTIFF: {image_path}")
                arr = ds.ReadAsArray()  # (bands, H, W) или (H, W)
                if arr.ndim == 2:
                    arr = arr[None, ...]
                height = ds.RasterYSize
                width = ds.RasterXSize
                gt = ds.GetGeoTransform()
                left = gt[0]
                top = gt[3]
                right = gt[0] + width * gt[1]
                bottom = gt[3] + height * gt[5]
                proj_wkt = ds.GetProjection()

                if max(height, width) > 3000:
                    step = max(1, math.ceil(max(height, width) / 3000))
                    arr = arr[:, ::step, ::step]
                    height = arr.shape[1]
                    width = arr.shape[2]

                if arr.shape[0] >= 3:
                    rgb = arr[:3].astype('float64')
                    rgb = (rgb - rgb.min()) / max(1e-9, (rgb.max() - rgb.min()))
                    rgb = rgb.transpose(1, 2, 0)
                    img_bg = rgb
                else:
                    gray = arr[0].astype('float64')
                    gray = (gray - gray.min()) / max(1e-9, (gray.max() - gray.min()))
                    img_bg = gray

                dest_crs = None
                if PYPROJ_AVAILABLE and proj_wkt:
                    try:
                        dest_crs = pyproj.CRS.from_wkt(proj_wkt)
                    except Exception:
                        dest_crs = None

                bg_tuple = (img_bg, (left, right, bottom, top), dest_crs)
            else:
                bg = plt.imread(image_path)
                bg_tuple = (bg, None, None)
        except Exception as e:
            print(f"[WARN] Не удалось загрузить подложку: {e}", flush=True)


    total = len(parcel_ids)

    def _process_one(cad_number: str, idx: int) -> None:
        print(f"[INFO] ({idx}/{total}) Формирую документ для КН {cad_number}...", flush=True)
        base_filename = f"parcel_{_safe_filename(cad_number)}"
        pdf_path = os.path.join(output_dir, base_filename + '.pdf')
        docx_path = os.path.join(output_dir, base_filename + '.docx')
        try:
            zoom_img_path: Optional[str] = None
            try:
                zoom_img_path = os.path.join(output_dir, f"zoom_{_safe_filename(cad_number)}.png")
                _make_zoom_image_for_parcel(
                    cad_number, cadastral_df, violations_df, viol_coords_df, coords_df,
                    zoom_img_path, background=bg_tuple, proj_string=proj_string,
                    min_point_spacing=min_point_spacing, max_points=max_points
                )
                if not os.path.exists(zoom_img_path):
                    zoom_img_path = None
            except Exception as e:
                print(f"[WARN] Не удалось создать zoom-изображение для {cad_number}: {e}", flush=True)

            overview_img_path: Optional[str] = None
            try:
                overview_img_path = os.path.join(output_dir, f"overview_{_safe_filename(cad_number)}.png")
                _make_overview_map(
                    cad_number, cadastral_df, violations_df, viol_coords_df,
                    overview_img_path, background=bg_tuple, proj_string=proj_string,
                    min_point_spacing=min_point_spacing, max_points=max_points
                )
                if not os.path.exists(overview_img_path):
                    overview_img_path = None
            except Exception as e:
                print(f"[WARN] Не удалось создать обзорную карту для {cad_number}: {e}", flush=True)

            story = build_parcel_story(
                cadastral_number=cad_number,
                cadastral_df=cadastral_df,
                coords_df=coords_df,
                violations_df=violations_df,
                viol_coords_df=viol_coords_df,
                image_path=zoom_img_path or image_path,
                overview_path=overview_img_path,
                font_name=font_name,
            )
            if out_format in ('pdf', 'both'):
                doc = SimpleDocTemplate(pdf_path, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
                doc.build(story)
            if out_format in ('docx', 'both'):
                if not DOCX_AVAILABLE:
                    raise RuntimeError("Требуется пакет 'python-docx' (pip install python-docx) для экспорта DOCX")
                _build_docx(
                    docx_path=docx_path,
                    cadastral_number=cad_number,
                    cadastral_df=cadastral_df,
                    coords_df=coords_df,
                    violations_df=violations_df,
                    viol_coords_df=viol_coords_df,
                    image_path=zoom_img_path or image_path,
                    overview_path=overview_img_path,
                )
        except KeyError as e:
            print(f"[WARN] Пропускаю {cad_number}: отсутствует столбец {e}", flush=True)
        except Exception as e:
            print(f"[WARN] Проблема при формировании {cad_number}: {e}", flush=True)

    max_workers = max(1, min(4, (os.cpu_count() or 1)))
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(_process_one, cad_number, idx): cad_number for idx, cad_number in enumerate(parcel_ids, 1)}
        for _ in as_completed(futures):
            pass

    print(f"[INFO] Готово. Файлы: {output_dir}", flush=True)


def _make_zoom_image_for_parcel(cadastral_number: str,
                                cadastral_df: pd.DataFrame,
                                violations_df: pd.DataFrame,
                                viol_coords_df: pd.DataFrame,
                                coords_df: pd.DataFrame,
                                save_path: str,
                                background: Optional[Tuple[Any, Optional[Tuple[float, float, float, float]], Any]] = None,
                                proj_string: Optional[str] = None,
                                min_point_spacing: float = 3.0,
                                max_points: Optional[int] = None) -> None:
    """
    Строит картинку с увеличением на нарушения конкретного участка.
    Источник подложки отсутствует в этом скрипте, поэтому рисуем схему:
    - контур участка (синий)
    - контуры нарушений для этого участка (красный)
    Масштаб – tight по объединённым границам.
    Берём WKT из листа "3. Нарушения" (если есть столбец 'WKT геометрии').
    """
    v = violations_df[violations_df['Ближайший кадастровый номер'] == cadastral_number]
    if v.empty:
        return

    # Попытка взять WKT геометрий нарушений
    geoms = []
    if 'WKT геометрии' in v.columns:
        for w in v['WKT геометрии'].dropna().astype(str).tolist():
            try:
                geom = shapely_wkt.loads(w)
                geoms.append(geom)
            except Exception:
                pass

    if not geoms:
        # Если нет WKT, попробуем приблизить по центроидам + площади (не точно)
        # В таком случае просто не строим картинку.
        return

    # Контур участка по WKT из листа 2 (если есть)
    parcel_geom = None
    if 'WKT геометрии' in cadastral_df.columns:
        cad_rows = cadastral_df[cadastral_df['Кадастровый номер'] == cadastral_number]
        if not cad_rows.empty:
            try:
                parcel_geom = shapely_wkt.loads(str(cad_rows.iloc[0]['WKT геометрии']))
            except Exception:
                parcel_geom = None

    fig, ax = plt.subplots(figsize=(5, 5), dpi=130)

    # --- Начало блока перепроецирования ---
    project = None
    dest_crs = None
    if background is not None and proj_string:
        if not PYPROJ_AVAILABLE:
            print("[WARN] Пакет 'pyproj' не найден. Пропускаю перепроецирование. Установите: pip install pyproj", flush=True)
        else:
            try:
                _, _, dest_crs = background
                if dest_crs:
                    print(f"[DEBUG] СК растра (GeoTIFF): {dest_crs.to_string()}", flush=True)
                    src_crs = pyproj.CRS.from_proj4(proj_string)
                    print(f"[DEBUG] Исходная СК векторов (proj-string): {src_crs.to_string()}", flush=True)
                    # always_xy=True важен для правильного порядка координат в shapely
                    transformer = pyproj.Transformer.from_crs(src_crs, dest_crs, always_xy=True)
                    project = transformer.transform
                    print("[DEBUG] Трансформация координат создана успешно.", flush=True)
            except Exception as e:
                print(f"[WARN] Не удалось создать трансформацию координат: {e}", flush=True)

    if project:
        try:
            geoms = [transform(project, g) for g in geoms]
            if parcel_geom:
                parcel_geom = transform(project, parcel_geom)
        except Exception as e:
            print(f"[WARN] Ошибка при перепроецировании геометрий: {e}", flush=True)
            # Если проекция не удалась, лучше не рисовать ничего, чем в неверных координатах
            return
    # --- Конец блока перепроецирования ---

    # Подложка: если заранее загружена — отрисуем как фон
    if background is not None:
        try:
            bg_img, bg_extent, _ = background
            if bg_img is not None:
                if bg_extent is not None:
                    left, right, bottom, top = bg_extent
                    ax.imshow(bg_img, extent=[left, right, bottom, top], alpha=0.8)
                else:
                    ax.imshow(bg_img)
        except Exception:
            pass

    if parcel_geom is not None and not parcel_geom.is_empty:
        _plot_shapely(ax, parcel_geom, facecolor='none', edgecolor='blue', linewidth=0.8, alpha=0.9)

    for g in geoms:
        _plot_shapely(ax, g, facecolor='red', edgecolor='darkred', linewidth=0.5, alpha=0.6)

    # Отображение центроидов нарушений
    if 'Центроид X' in v.columns and 'Центроид Y' in v.columns:
        for _, row in v.iterrows():
            try:
                cx = pd.to_numeric(row['Центроид X'], errors='coerce')
                cy = pd.to_numeric(row['Центроид Y'], errors='coerce')
                if pd.notna(cx) and pd.notna(cy):
                    if project:
                        cx, cy = project(cx, cy)
                    ax.plot(cx, cy, marker='+', markersize=6, color='yellow', markeredgewidth=1.5)
            except Exception:
                pass

    # Нумерация точек по КАЖДОМУ нарушению этого КН (лист 5)
    vs_all = violations_df[violations_df['Ближайший кадастровый номер'] == cadastral_number]
    if not vs_all.empty and not viol_coords_df.empty and '№ нарушения' in viol_coords_df.columns:
        try:
            vs = vs_all.reset_index(drop=True)
            for local_idx, (idx_row, row_v) in enumerate(vs.iterrows(), 1):
                # Используем глобальный номер для поиска координат
                global_v_num = row_v['№ нарушения'] if '№ нарушения' in vs_all.columns else (idx_row + 1)
                try:
                    global_v_num = int(global_v_num)
                except Exception:
                    global_v_num = idx_row + 1
                
                sub = viol_coords_df[viol_coords_df['№ нарушения'] == global_v_num].copy()
                if sub.empty:
                    continue
                sub.sort_values(by='Номер точки', inplace=True)
                pts = list(zip(pd.to_numeric(sub['X'], errors='coerce').astype(float),
                                pd.to_numeric(sub['Y'], errors='coerce').astype(float)))
                if project and pts:
                    try:
                        pts = [project(x, y) for (x, y) in pts]
                    except Exception as e:
                        print(f"[WARN] Не удалось перепроецировать точки нарушения {global_v_num}: {e}", flush=True)
                        continue

                if len(pts) >= 2 and pts[0] == pts[-1]:
                    pts = pts[:-1]
                
                pts_simplified = _simplify_points(pts, min_spacing=min_point_spacing, max_points=max_points)
                print(f"[DEBUG] Image: КН {cadastral_number}, Локальное нарушение №{local_idx} (глобальное #{global_v_num}): {len(pts)} точек -> {len(pts_simplified)} после упрощения.", flush=True)
                
                for i, (x, y) in enumerate(pts_simplified, 1):
                    ax.plot([x], [y], marker='o', markersize=2, color='darkred')
                    ax.text(x, y, str(i), fontsize=5, color='darkred', ha='center', va='bottom',
                            bbox=dict(boxstyle='round,pad=0.1', facecolor='white', edgecolor='none', alpha=0.8))
        except Exception:
            pass

    # Масштаб по объединённым границам
    all_bounds = [g.bounds for g in geoms if g and not g.is_empty]
    if parcel_geom is not None and not parcel_geom.is_empty:
        all_bounds.append(parcel_geom.bounds)

    if not all_bounds:
        print(f"[WARN] Нет геометрий для определения масштаба для КН {cadastral_number}", flush=True)
    else:
        minx = min(b[0] for b in all_bounds)
        miny = min(b[1] for b in all_bounds)
        maxx = max(b[2] for b in all_bounds)
        maxy = max(b[3] for b in all_bounds)
        vector_bbox_tuple = (minx, miny, maxx, maxy)

        final_minx, final_miny, final_maxx, final_maxy = vector_bbox_tuple

        # Проверка пересечения с подложкой
        if background is not None and background[1] is not None:
            bg_extent = background[1]
            # shapely.box: minx, miny, maxx, maxy | raster bounds: left, bottom, right, top
            raster_bbox_tuple = (bg_extent[0], bg_extent[2], bg_extent[1], bg_extent[3])

            if project:
                vector_box = box(*vector_bbox_tuple)
                raster_box = box(*raster_bbox_tuple)

                if not vector_box.intersects(raster_box):
                    print("[WARN] ВНИМАНИЕ: Перепроецированные данные НЕ ПЕРЕСЕКАЮТСЯ с подложкой GeoTIFF.", flush=True)
                    print(f"[DEBUG] Границы векторов (в СК растра): {vector_bbox_tuple}", flush=True)
                    print(f"[DEBUG] Границы растра: {raster_bbox_tuple}", flush=True)
                    print("[INFO] Оставляю масштаб по векторам для сохранения зума.", flush=True)

        padding_x = (final_maxx - final_minx) * 1.0 if final_maxx > final_minx else 50
        padding_y = (final_maxy - final_miny) * 1.0 if final_maxy > final_miny else 50
        ax.set_xlim(final_minx - padding_x, final_maxx + padding_x)
        ax.set_ylim(final_miny - padding_y, final_maxy + padding_y)

    ax.set_aspect('equal', adjustable='box')
    ax.axis('off')

    # Жёсткое и простое ограничение размера, чтобы гарантированно не вылетать по "Image size ... too large"
    max_pixels = 65536
    safe_figsize = (4.0, 4.0)  # дюймы
    safe_dpi = 200
    fig.set_size_inches(*safe_figsize)

    with warnings.catch_warnings():
        warnings.simplefilter("ignore", UserWarning)
        try:
            plt.tight_layout(pad=0.5)
        except Exception:
            pass

    try:
        if PIL_AVAILABLE:
            buf = BytesIO()
            fig.savefig(buf, format='png', dpi=safe_dpi)
            buf.seek(0)

            img = PILImage.open(buf)
            width, height = img.size

            if width > max_pixels or height > max_pixels:
                scale = min(max_pixels / width, max_pixels / height)
                new_size = (max(1, int(width * scale)), max(1, int(height * scale)))
                print(f"[DEBUG] Pillow zoom: уменьшаю {width}x{height} до {new_size[0]}x{new_size[1]}", flush=True)
                try:
                    img = img.resize(new_size, PILImage.Resampling.LANCZOS)
                except AttributeError:
                    img = img.resize(new_size, PILImage.LANCZOS)

            img.save(save_path, 'PNG', optimize=True)
            buf.close()
        else:
            fig.savefig(save_path, dpi=safe_dpi)
    except Exception as e:
        # Не пробрасываем ошибку наверх, чтобы экспорт по участкам не прерывался
        print(f"[WARN] Не удалось сохранить zoom-изображение для {cadastral_number}: {e}", flush=True)
    finally:
        plt.close(fig)


def _plot_shapely(ax, geom, facecolor='none', edgecolor='black', linewidth=1.0, alpha=1.0):
    if geom is None or geom.is_empty:
        return
    if isinstance(geom, (MultiPolygon,)):
        for g in geom.geoms:
            _plot_shapely(ax, g, facecolor=facecolor, edgecolor=edgecolor, linewidth=linewidth, alpha=alpha)
        return
    if isinstance(geom, Polygon):
        exterior = np.asarray(geom.exterior.coords)
        verts = exterior.tolist()
        codes = [MplPath.MOVETO] + [MplPath.LINETO] * (len(exterior) - 2) + [MplPath.CLOSEPOLY]

        for interior in geom.interiors:
            ring = np.asarray(interior.coords)
            verts.extend(ring.tolist())
            codes.extend([MplPath.MOVETO] + [MplPath.LINETO] * (len(ring) - 2) + [MplPath.CLOSEPOLY])

        path = MplPath(verts, codes)
        patch = PathPatch(path, facecolor=facecolor, edgecolor=edgecolor, linewidth=linewidth, alpha=alpha)
        ax.add_patch(patch)


def _make_overview_map(cadastral_number: str,
                       cadastral_df: pd.DataFrame,
                       violations_df: pd.DataFrame,
                       viol_coords_df: pd.DataFrame,
                       save_path: str,
                       background: Optional[Tuple[Any, Optional[Tuple[float, float, float, float]], Any]] = None,
                       proj_string: Optional[str] = None,
                       min_point_spacing: float = 3.0,
                       max_points: Optional[int] = None) -> None:
    v = violations_df[violations_df['Ближайший кадастровый номер'] == cadastral_number]
    
    cad_rows = cadastral_df[cadastral_df['Кадастровый номер'] == cadastral_number]
    if cad_rows.empty:
        return
    
    cad_row = cad_rows.iloc[0]
    
    parcel_geom = None
    if 'WKT геометрии' in cadastral_df.columns:
        try:
            parcel_geom = shapely_wkt.loads(str(cad_row['WKT геометрии']))
        except Exception:
            pass
    
    if parcel_geom is None or parcel_geom.is_empty:
        return
    
    geoms = []
    if not v.empty and 'WKT геометрии' in v.columns:
        for idx, w in enumerate(v['WKT геометрии'].dropna().astype(str).tolist(), 1):
            try:
                geom = shapely_wkt.loads(w)
                geoms.append(geom)
                if parcel_geom and geom:
                    intersects = parcel_geom.intersects(geom)
                    if intersects:
                        intersection_area = parcel_geom.intersection(geom).area
                        ratio = intersection_area / geom.area if geom.area > 0 else 0
                        print(f"[DEBUG] КН {cadastral_number}, Нарушение {idx}: пересечение {intersection_area:.2f} м² ({ratio*100:.1f}%)")
                    else:
                        distance = parcel_geom.distance(geom)
                        print(f"[DEBUG] КН {cadastral_number}, Нарушение {idx}: НЕТ пересечения, расстояние {distance:.2f} м")
            except Exception as e:
                print(f"[WARN] Ошибка загрузки геометрии нарушения: {e}")
                pass
    
    area_of_interest = parcel_geom
    if geoms:
        from shapely.ops import unary_union
        all_geoms_combined = unary_union([parcel_geom] + geoms)
        buffer_distance = max(parcel_geom.bounds[2] - parcel_geom.bounds[0],
                              parcel_geom.bounds[3] - parcel_geom.bounds[1]) * 0.3
        area_of_interest = all_geoms_combined.buffer(buffer_distance)
    
    all_cadastral_info = []
    if 'WKT геометрии' in cadastral_df.columns:
        for idx, row in cadastral_df.iterrows():
            try:
                cad_num = row.get('Кадастровый номер', '')
                if pd.isna(cad_num):
                    continue
                geom = shapely_wkt.loads(str(row['WKT геометрии']))
                if geom and not geom.is_empty:
                    is_main = (str(cad_num) == str(cadastral_number))
                    if is_main or geom.intersects(area_of_interest) or geom.distance(area_of_interest) < buffer_distance * 0.5:
                        all_cadastral_info.append({
                            'number': str(cad_num),
                            'geom': geom,
                            'is_main': is_main
                        })
            except Exception:
                pass
    
    fig, ax = plt.subplots(figsize=(8, 8), dpi=150)
    
    project = None
    dest_crs = None
    if background is not None and proj_string:
        if PYPROJ_AVAILABLE:
            try:
                _, _, dest_crs = background
                if dest_crs:
                    src_crs = pyproj.CRS.from_proj4(proj_string)
                    transformer = pyproj.Transformer.from_crs(src_crs, dest_crs, always_xy=True)
                    project = transformer.transform
            except Exception:
                pass
    
    if project:
        try:
            parcel_geom = transform(project, parcel_geom)
            geoms = [transform(project, g) for g in geoms]
            for info in all_cadastral_info:
                info['geom'] = transform(project, info['geom'])
        except Exception:
            return
    
    if background is not None:
        try:
            bg_img, bg_extent, _ = background
            if bg_img is not None and bg_extent is not None:
                left, right, bottom, top = bg_extent
                ax.imshow(bg_img, extent=[left, right, bottom, top], alpha=0.7)
        except Exception:
            pass
    
    for info in all_cadastral_info:
        if info['is_main']:
            _plot_shapely(ax, info['geom'], facecolor='lightblue', edgecolor='blue', linewidth=2.0, alpha=0.4)
        else:
            _plot_shapely(ax, info['geom'], facecolor='lightyellow', edgecolor='orange', linewidth=1.5, alpha=0.3)
    
    cadastral_area = cad_row.get('Площадь, м²', 0)
    try:
        cadastral_area = float(pd.to_numeric(cadastral_area, errors='coerce'))
    except Exception:
        cadastral_area = 0
    
    total_violation_area = 0
    if not v.empty:
        total_violation_area = float(v['Площадь нарушения, м²'].fillna(0).sum())
    
    centroid = parcel_geom.centroid
    ax.text(centroid.x, centroid.y, 
            f"КН: {cadastral_number}\nПлощадь: {cadastral_area:.2f} м²\nНарушений: {total_violation_area:.2f} м²",
            fontsize=8, color='blue', ha='center', va='center',
            bbox=dict(boxstyle='round,pad=0.5', facecolor='white', edgecolor='blue', alpha=0.9))
    
    for g in geoms:
        _plot_shapely(ax, g, facecolor='red', edgecolor='darkred', linewidth=1.0, alpha=0.5)
    
    vs_all = violations_df[violations_df['Ближайший кадастровый номер'] == cadastral_number]
    if not vs_all.empty and not viol_coords_df.empty and '№ нарушения' in viol_coords_df.columns:
        try:
            vs = vs_all.reset_index(drop=True)
            for local_idx, (idx_row, row_v) in enumerate(vs.iterrows(), 1):
                global_v_num = row_v['№ нарушения'] if '№ нарушения' in vs_all.columns else (idx_row + 1)
                try:
                    global_v_num = int(global_v_num)
                except Exception:
                    global_v_num = idx_row + 1
                
                sub = viol_coords_df[viol_coords_df['№ нарушения'] == global_v_num].copy()
                if sub.empty:
                    continue
                sub.sort_values(by='Номер точки', inplace=True)
                pts = list(zip(pd.to_numeric(sub['X'], errors='coerce').astype(float),
                                pd.to_numeric(sub['Y'], errors='coerce').astype(float)))
                
                if project and pts:
                    try:
                        pts = [project(x, y) for (x, y) in pts]
                    except Exception:
                        continue
                
                if len(pts) >= 2 and pts[0] == pts[-1]:
                    pts = pts[:-1]
                
                pts_simplified = _simplify_points(pts, min_spacing=min_point_spacing, max_points=max_points)
                
                violation_area = row_v.get('Площадь нарушения, м²', 0)
                try:
                    violation_area = float(pd.to_numeric(violation_area, errors='coerce'))
                except Exception:
                    violation_area = 0
                
                if pts_simplified:
                    centroid_x = sum(p[0] for p in pts_simplified) / len(pts_simplified)
                    centroid_y = sum(p[1] for p in pts_simplified) / len(pts_simplified)
                    ax.text(centroid_x, centroid_y, 
                            f"№{local_idx}\n{violation_area:.1f} м²",
                            fontsize=7, color='white', ha='center', va='center', weight='bold',
                            bbox=dict(boxstyle='round,pad=0.3', facecolor='red', edgecolor='darkred', alpha=0.8))
                
                for i, (x, y) in enumerate(pts_simplified, 1):
                    ax.plot([x], [y], marker='o', markersize=2.5, color='yellow', markeredgecolor='darkred', markeredgewidth=0.5)
                    ax.text(x, y, str(i), fontsize=5, color='black', ha='center', va='center', weight='bold')
        except Exception:
            pass
    
    all_bounds = []
    for info in all_cadastral_info:
        if info['geom'] and not info['geom'].is_empty:
            all_bounds.append(info['geom'].bounds)
    for g in geoms:
        if g and not g.is_empty:
            all_bounds.append(g.bounds)
    
    if all_bounds:
        minx = min(b[0] for b in all_bounds)
        miny = min(b[1] for b in all_bounds)
        maxx = max(b[2] for b in all_bounds)
        maxy = max(b[3] for b in all_bounds)
        
        padding_x = (maxx - minx) * 0.2
        padding_y = (maxy - miny) * 0.2
        ax.set_xlim(minx - padding_x, maxx + padding_x)
        ax.set_ylim(miny - padding_y, maxy + padding_y)
    
    ax.set_aspect('equal', adjustable='box')
    ax.set_title(f'Обзорная карта участка {cadastral_number}', fontsize=10, weight='bold')
    ax.grid(True, alpha=0.3, linestyle='--', linewidth=0.5)

    # Простое ограничение размера обзорной карты, чтобы избежать ошибок размера изображения
    max_pixels = 65536
    safe_figsize = (6.0, 6.0)
    safe_dpi = 200
    fig.set_size_inches(*safe_figsize)

    with warnings.catch_warnings():
        warnings.simplefilter("ignore", UserWarning)
        try:
            plt.tight_layout(pad=0.7)
        except Exception:
            pass

    try:
        if PIL_AVAILABLE:
            buf = BytesIO()
            fig.savefig(buf, format='png', dpi=safe_dpi)
            buf.seek(0)

            img = PILImage.open(buf)
            width, height = img.size

            if width > max_pixels or height > max_pixels:
                scale = min(max_pixels / width, max_pixels / height)
                new_size = (max(1, int(width * scale)), max(1, int(height * scale)))
                print(f"[DEBUG] Pillow overview: уменьшаю {width}x{height} до {new_size[0]}x{new_size[1]}", flush=True)
                try:
                    img = img.resize(new_size, PILImage.Resampling.LANCZOS)
                except AttributeError:
                    img = img.resize(new_size, PILImage.LANCZOS)

            img.save(save_path, 'PNG', optimize=True)
            buf.close()
        else:
            fig.savefig(save_path, dpi=safe_dpi)
    except Exception as e:
        print(f"[WARN] Не удалось сохранить обзорную карту для {cadastral_number}: {e}", flush=True)
    finally:
        plt.close(fig)


def _build_docx(docx_path: str,
                cadastral_number: str,
                cadastral_df: pd.DataFrame,
                coords_df: pd.DataFrame,
                violations_df: pd.DataFrame,
                viol_coords_df: pd.DataFrame,
                image_path: Optional[str],
                overview_path: Optional[str] = None) -> None:
    doc = Document()

    # Настройка стилей для кириллицы
    for style_name in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3']:
        try:
            st = doc.styles[style_name]
            st.font.name = 'Times New Roman'
            st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            st.font.size = Pt(11)
        except Exception:
            pass

    h = doc.add_paragraph('СХЕМА')
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph('обмера земельного участка')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Данные по участку
    doc.add_paragraph(f"Кадастровый номер: {cadastral_number}")

    # Сводка по нарушениям
    v = violations_df[violations_df['Ближайший кадастровый номер'] == cadastral_number]
    if not v.empty:
        sum_violation = float(v['Площадь нарушения, м²'].fillna(0).sum())
        doc.add_paragraph(f"Количество нарушений: {len(v)}")
        doc.add_paragraph(f"Суммарная площадь нарушений: {sum_violation:.2f} м²")

        v_table = doc.add_table(rows=1, cols=4)
        try:
            v_table.style = 'Table Grid'
        except Exception:
            pass
        v_hdr = v_table.rows[0].cells
        v_hdr[0].text = '№'
        v_hdr[1].text = 'Площадь, м²'
        v_hdr[2].text = 'Центроид X'
        v_hdr[3].text = 'Центроид Y'

        for i, (_, r) in enumerate(v.iterrows(), 1):
            try:
                area_val = float(pd.to_numeric(r.get('Площадь нарушения, м²', 0), errors='coerce'))
            except Exception:
                area_val = 0
            try:
                cx = float(pd.to_numeric(r.get('Центроид X', 0), errors='coerce'))
            except Exception:
                cx = 0
            try:
                cy = float(pd.to_numeric(r.get('Центроид Y', 0), errors='coerce'))
            except Exception:
                cy = 0
            px, py = _present_xy(cx, cy)

            row = v_table.add_row().cells
            row[0].text = str(i)
            row[1].text = f"{area_val:.2f}"
            row[2].text = f"{px:.2f}"
            row[3].text = f"{py:.2f}"

        # Таблицы координат по каждому нарушению (лист 5)
        if DOCX_AVAILABLE and not viol_coords_df.empty and '№ нарушения' in viol_coords_df.columns:
            v_reset = v.reset_index(drop=True)
            for local_idx, (idx_row, row_v) in enumerate(v_reset.iterrows(), 1):
                # Используем глобальный номер для поиска координат
                global_v_num = row_v['№ нарушения'] if '№ нарушения' in v.columns else (idx_row + 1)
                try:
                    global_v_num = int(global_v_num)
                except Exception:
                    global_v_num = idx_row + 1
                
                sub = viol_coords_df[viol_coords_df['№ нарушения'] == global_v_num].copy()
                if sub.empty:
                    continue
                sub.sort_values(by='Номер точки', inplace=True)
                pts = list(zip(pd.to_numeric(sub['X'], errors='coerce').astype(float),
                                pd.to_numeric(sub['Y'], errors='coerce').astype(float)))
                if len(pts) >= 2 and pts[0] == pts[-1]:
                    pts = pts[:-1]
                print(f"[DEBUG] DOCX: КН {cadastral_number}, Локальное нарушение №{local_idx} (глобальное #{global_v_num}): {len(pts)} точек в таблице.", flush=True)
                dists = compute_distances(pts)

                # Используем локальный номер в заголовке
                doc.add_paragraph(f"Координаты нарушения № {local_idx}")
                vt = doc.add_table(rows=1, cols=4)
                try:
                    vt.style = 'Table Grid'
                except Exception:
                    pass
                hdr = vt.rows[0].cells
                hdr[0].text = 'Обозначение\nточки'
                hdr[1].text = 'X'
                hdr[2].text = 'Y'
                hdr[3].text = 'Расстояние до точки, м.'
                for j, (x, y) in enumerate(pts, 1):
                    dist_idx = (j - 2) % len(dists)
                    px, py = _present_xy(x, y)
                    rw = vt.add_row().cells
                    rw[0].text = str(j)
                    rw[1].text = f"{px:.2f}"
                    rw[2].text = f"{py:.2f}"
                    rw[3].text = f"{dists[dist_idx]:.2f}"

    if overview_path and os.path.exists(overview_path):
        try:
            doc.add_paragraph('Обзорная карта участка')
            doc.add_picture(overview_path, width=Inches(5.5))
        except Exception:
            pass

    if image_path and os.path.exists(image_path):
        try:
            doc.add_paragraph('Детальная карта нарушений')
            doc.add_picture(image_path, width=Inches(5.5))
        except Exception:
            pass

    doc.save(docx_path)


def main() -> None:
    parser = argparse.ArgumentParser(description='Экспорт отдельных документов по каждому участку (PDF/DOCX).')
    parser.add_argument('--report', default=os.path.join('output', 'comprehensive', 'report.xlsx'), help='Путь к итоговому report.xlsx')
    parser.add_argument('--output', default=os.path.join('output', 'comprehensive', 'per_parcel_docs'), help='Папка для документов')
    parser.add_argument('--image', default=os.path.join('output', 'comprehensive', 'visualization.png'), help='Картинка (например visualization.png) для вставки в документ')
    parser.add_argument('--all', action='store_true', help='Экспорт для всех участков, а не только имеющих нарушения')
    parser.add_argument('--limit', type=int, default=None, help='Ограничить число обрабатываемых участков (для проверки)')
    parser.add_argument('--font', type=str, default=None, help='Путь к TTF-шрифту с поддержкой кириллицы (например C:\\Windows\\Fonts\\arial.ttf)')
    parser.add_argument('--format', type=str, choices=['pdf', 'docx', 'both'], default='pdf', help='Формат выходных файлов')
    parser.add_argument('--geotiff', type=str, default=os.path.join('geotiffs', 'input.tiff'), help='Путь к GeoTIFF для подложки (фоновая карта)')
    parser.add_argument('--proj-string', type=str,
                        default='+proj=tmerc +lat_0=0 +lon_0=109.03333333333 +k=1 +x_0=4250000 +y_0=-5211057.63 +ellps=krass +towgs84=23.57,-140.95,-79.8,0,0.35,0.79,-0.22 +units=m +no_defs',
                        help='PROJ-строка для исходной системы координат векторов (например, МСК)')
    parser.add_argument('--min-point-spacing', type=float, default=3.0,
                        help='Минимальное расстояние между точками нарушений (м) для упрощения визуализации')
    parser.add_argument('--max-points', type=int, default=None,
                        help='Максимальное количество точек для отображения на одно нарушение')
    args = parser.parse_args()

    if not os.path.exists(args.report):
        print(f"[ERROR] Не найден файл отчёта: {args.report}", flush=True)
        sys.exit(1)

    try:
        export_per_parcel_pdfs(
            report_path=args.report,
            output_dir=args.output,
            image_path=args.geotiff if args.geotiff and os.path.exists(args.geotiff) else (args.image if args.image and os.path.exists(args.image) else None),
            only_with_violations=not args.all,
            limit=args.limit,
            font_path=args.font,
            out_format=args.format,
            proj_string=args.proj_string,
            min_point_spacing=args.min_point_spacing,
            max_points=args.max_points,
        )
    except Exception as e:
        print(f"[ERROR] Экспорт прерван: {e}", flush=True)
        sys.exit(2)


if __name__ == '__main__':
    main()


